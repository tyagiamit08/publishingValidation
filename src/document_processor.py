from docx import Document
import PyPDF2
import logging
import pdfplumber
import io

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def process_docx(doc_path):
    """Process a DOCX file and extract its text content."""
    try:
        logging.info(f"Attempting to process DOCX file: {doc_path}")
        doc = Document(doc_path)
        full_text = []
        for para in doc.paragraphs:
            logging.debug(f"Extracted paragraph: {para.text[:50]}...")  # Log the first 50 characters of each paragraph
            full_text.append(para.text)
        logging.info(f"Successfully processed DOCX file: {doc_path}")
        return "\n".join(full_text)
    except Exception as e:
        logging.error(f"Error reading DOCX file: {str(e)}", exc_info=True)
        return f"Error reading DOCX file: {str(e)}"

def process_pdf(doc_path):
    """Process a PDF file and extract its text content."""
    try:
        with open(doc_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            full_text = []
            for page in reader.pages:
                full_text.append(page.extract_text())
            return "\n".join(full_text)
    except Exception as e:
        logging.error(f"Error reading PDF file: {str(e)}", exc_info=True)
        return f"Error reading PDF file: {str(e)}"

def extract_images_from_pdf(file_bytes):
    images = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            im = page.to_image(resolution=300)
            img_bytes = io.BytesIO()
            im.original.save(img_bytes, format='PNG')
            images.append(img_bytes.getvalue())
    return images

def extract_images_from_docx(file_bytes):
    images = []
    doc = Document(io.BytesIO(file_bytes))
    for rel in doc.part._rels:
        rel = doc.part._rels[rel]
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            images.append(image_data)
    return images
